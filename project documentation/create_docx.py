"""Convert CLIENT_WEBSITE_SPECIFICATION.md to DOCX"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_paragraph_with_format(doc, text, style=None, bold=False):
    """Add paragraph, handling bold/italic inline formatting"""
    if not text.strip():
        return
    p = doc.add_paragraph(style=style)
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2] + ' ')
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1] + ' ')
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1] + ' ')
            run.font.name = 'Consolas'
        else:
            p.add_run(part)
    return p

def parse_and_build(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i]
        line_stripped = line.rstrip()
        
        # Skip empty lines at start
        if not line_stripped and not in_table:
            i += 1
            continue
            
        # Horizontal rule
        if line_stripped == '---':
            doc.add_paragraph()
            i += 1
            continue
            
        # Headings
        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=0)
            i += 1
            continue
        if line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=1)
            i += 1
            continue
        if line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=2)
            i += 1
            continue
            
        # Table - detect table rows
        if '|' in line_stripped and line_stripped.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            # Skip separator row
            if re.match(r'^\|[\s\-:|]+\|$', line_stripped):
                i += 1
                continue
            cells = [c.strip() for c in line_stripped.split('|')[1:-1]]
            if cells:
                table_data.append(cells)
            i += 1
            continue
        else:
            if in_table and table_data:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            table.rows[row_idx].cells[col_idx].text = cell_text.replace('`', '')
                doc.add_paragraph()
                table_data = []
            in_table = False
            
        # Code block
        if line_stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            if i < len(lines):
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.add_run('\n'.join(code_lines)).font.name = 'Consolas'
                p.paragraph_format.left_indent = Inches(0.5)
            continue
            
        # Bullet list
        if line_stripped.startswith('- ') or line_stripped.startswith('* '):
            text = line_stripped[2:].strip()
            add_paragraph_with_format(doc, text, style='List Bullet')
            i += 1
            continue
            
        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)$', line_stripped)
        if num_match:
            text = num_match.group(2).strip()
            add_paragraph_with_format(doc, text, style='List Number')
            i += 1
            continue
            
        # Checkbox
        if line_stripped.strip().startswith('- [ ]'):
            text = line_stripped.strip()[6:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.add_run('☐ ' + text)
            i += 1
            continue
            
        # Regular paragraph
        if line_stripped:
            add_paragraph_with_format(doc, line_stripped)
            
        i += 1
    
    # Handle any remaining table
    if table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    table.rows[row_idx].cells[col_idx].text = cell_text.replace('`', '')
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == '__main__':
    parse_and_build(
        r'G:\Samprojects\Old house\CLIENT_WEBSITE_SPECIFICATION.md',
        r'G:\Samprojects\Old house\CLIENT_WEBSITE_SPECIFICATION.docx'
    )
